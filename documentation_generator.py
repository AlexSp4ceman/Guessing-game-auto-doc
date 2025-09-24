"""
DOCUMENTATION GENERATOR
=======================

This module automatically extracts comments from Python files and
generates a Word document with comprehensive documentation.

Dependencies: python-docx
"""

import ast
import re
import os
from docx import Document
from docx.shared import Inches
from datetime import datetime


class DocumentationGenerator:
    """
    A class to generate Word documentation from Python code comments.
    
    This class parses Python files, extracts comments and docstrings,
    and creates a formatted Word document with the documentation.
    """
    
    def __init__(self):
        """Initialize the documentation generator."""
        self.doc = Document()
        self.comments_data = []
    
    def cleanup_old_documentation(self, output_file):
        """
        Delete old documentation file if it exists.
        
        Args:
            output_file (str): Path to the documentation file to delete
        """
        if os.path.exists(output_file):
            try:
                os.remove(output_file)
                print(f"Deleted old documentation: {output_file}")
            except Exception as e:
                print(f"Warning: Could not delete old documentation: {e}")
        else:
            print("No old documentation found. Creating new file.")
    
    def parse_python_file(self, file_path):
        """
        Parse a Python file and extract all comments and docstrings.
        
        Args:
            file_path (str): Path to the Python file to parse
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Extract regular comments
            comment_pattern = r'#(.+)'
            comments = re.findall(comment_pattern, content)
            
            # Extract docstrings using AST
            tree = ast.parse(content)
            docstrings = []
            
            for node in ast.walk(tree):
                if isinstance(node, (ast.FunctionDef, ast.ClassDef, ast.Module)):
                    docstring = ast.get_docstring(node)
                    if docstring:
                        docstrings.append({
                            'type': type(node).__name__,
                            'name': getattr(node, 'name', 'Module'),
                            'docstring': docstring
                        })
            
            self.comments_data.append({
                'file_path': file_path,
                'comments': comments,
                'docstrings': docstrings
            })
            
        except Exception as e:
            print(f"Error parsing file {file_path}: {e}")
    
    def create_title_page(self):
        """Create the title page for the documentation."""
        title = self.doc.add_heading('Python Guessing Game Documentation', 0)
        title.alignment = 1  # Center alignment
        
        self.doc.add_paragraph()
        
        # Add metadata
        current_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        meta = self.doc.add_paragraph()
        meta.add_run("Generated on: ").bold = True
        meta.add_run(current_date)
        
        meta.add_run("\nAuthor: ").bold = True
        meta.add_run("Python Developer")
        
        meta.add_run("\nVersion: ").bold = True
        meta.add_run("1.0")
        
        self.doc.add_page_break()
    
    def create_table_of_contents(self):
        """Create a table of contents for the documentation."""
        heading = self.doc.add_heading('Table of Contents', 1)
        
        toc = self.doc.add_paragraph()
        toc.add_run("1. Introduction\n").bold = True
        toc.add_run("2. File Documentation\n").bold = True
        toc.add_run("3. Class and Function Documentation\n").bold = True
        toc.add_run("4. Code Comments\n").bold = True
        
        self.doc.add_page_break()
    
    def create_introduction(self):
        """Create the introduction section."""
        heading = self.doc.add_heading('1. Introduction', 1)
        
        intro = self.doc.add_paragraph()
        intro.add_run("This document contains automatically generated documentation ")
        intro.add_run("for the Python Guessing Game application. The documentation ")
        intro.add_run("includes code comments, docstrings, and function descriptions ")
        intro.add_run("extracted directly from the source code.")
        
        self.doc.add_paragraph()
        
        features = self.doc.add_paragraph("Key Features:")
        features.add_run("\n• Number guessing game with range 1-5")
        features.add_run("\n• Statistics tracking")
        features.add_run("\n• Input validation")
        features.add_run("\n• Automated documentation generation")
    
    def create_file_documentation(self):
        """Create documentation for each parsed file."""
        heading = self.doc.add_heading('2. File Documentation', 1)
        
        for file_data in self.comments_data:
            file_heading = self.doc.add_heading(f'File: {file_data["file_path"]}', 2)
            
            if file_data['docstrings']:
                for docstring in file_data['docstrings']:
                    if docstring['type'] == 'Module':
                        module_doc = self.doc.add_paragraph()
                        module_doc.add_run("Module Documentation:").bold = True
                        module_doc.add_run(f"\n{docstring['docstring']}")
    
    def create_class_function_docs(self):
        """Create documentation for classes and functions."""
        heading = self.doc.add_heading('3. Class and Function Documentation', 1)
        
        for file_data in self.comments_data:
            if file_data['docstrings']:
                for docstring in file_data['docstrings']:
                    if docstring['type'] in ['ClassDef', 'FunctionDef']:
                        doc_heading = self.doc.add_heading(
                            f"{docstring['type']}: {docstring['name']}", 3
                        )
                        
                        doc_paragraph = self.doc.add_paragraph()
                        doc_paragraph.add_run(docstring['docstring'])
                        
                        self.doc.add_paragraph()
    
    def create_comments_section(self):
        """Create a section with all extracted comments."""
        heading = self.doc.add_heading('4. Code Comments', 1)
        
        for file_data in self.comments_data:
            file_heading = self.doc.add_heading(f'Comments from {file_data["file_path"]}', 2)
            
            if file_data['comments']:
                for i, comment in enumerate(file_data['comments'], 1):
                    comment_para = self.doc.add_paragraph()
                    comment_para.add_run(f"Comment {i}: ").bold = True
                    comment_para.add_run(comment.strip())
            else:
                self.doc.add_paragraph("No comments found in this file.")
    
    def generate_documentation(self, output_file='Guessing_Game_Documentation.docx'):
        """
        Generate the complete Word documentation.
        
        Args:
            output_file (str): Name of the output Word document
        """
        print("Starting documentation generation...")
        
        # Clean up old documentation
        self.cleanup_old_documentation(output_file)
        
        # Create document sections
        self.create_title_page()
        self.create_table_of_contents()
        self.create_introduction()
        self.create_file_documentation()
        self.create_class_function_docs()
        self.create_comments_section()
        
        # Save document
        self.doc.save(output_file)
        print(f"Documentation generated successfully: {output_file}")
        print(f"Total files processed: {len(self.comments_data)}")
        
        # Show statistics
        total_comments = sum(len(file_data['comments']) for file_data in self.comments_data)
        total_docstrings = sum(len(file_data['docstrings']) for file_data in self.comments_data)
        print(f"Total comments extracted: {total_comments}")
        print(f"Total docstrings extracted: {total_docstrings}")


def main():
    """Main function to generate documentation for the guessing game."""
    generator = DocumentationGenerator()
    
    # Parse the Python files
    generator.parse_python_file('guessing_game.py')
    generator.parse_python_file('documentation_generator.py')
    
    # Generate the documentation
    generator.generate_documentation('Guessing_Game_Documentation.docx')


if __name__ == "__main__":
    main()